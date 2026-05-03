"""
Build English Word .docx from ver3.5_F-rev2 manuscript.
Embeds Tables 1-4 + Figures 1-3 + Supp Fig S1.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("C:/Projects/legend-team/reports/2026-05-02_copd-paper-part2-residual")
SRC = ROOT / "full_manuscript_ver3.5_F-rev1_E.md"
OUT = ROOT / "Manuscript_ver3.5_F-rev2_E.docx"

IMG = {
    "Figure 1": ("C:/Projects/COPD/Fig. 1 (2).jpg", "Conceptual directed acyclic graph (DAG)"),
    "Figure 2": ("C:/Projects/COPD/Fig. 2..png", "Cohort flow diagram"),
    "Figure 3": ("C:/Projects/COPD/Data/ver2.0/submission_session174_2026-05-02/02_analysis_outputs/Fig3_forest_NIE.png", "Forest plot of NIE"),
    "Supp S1": ("C:/Projects/COPD/Data/ver2.0/submission_session174_2026-05-02/02_analysis_outputs/S1_time_ordering.png", "Time-ordering of measurements"),
}

# ---- Document setup ----
doc = Document()
sec = doc.sections[0]
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5);  sec.bottom_margin = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_korean_font(run, name='Malgun Gothic'):
    """Set East Asian font for a run (needed for mixed-language docs)."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1f, 0x40, 0x68)
    return h

def add_para(text, italic=False, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    if italic: r.italic = True
    if bold: r.bold = True
    if size: r.font.size = Pt(size)
    return p

def add_md_para(md_text):
    """Convert markdown-flavored paragraph (with **bold**, *italic*, [refs], sub/sup) to docx."""
    p = doc.add_paragraph()
    # Split on bold/italic markers, preserve order
    # Simple approach: tokenize **...**, *...*, plain
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|<sub>[^<]+</sub>|<sup>[^<]+</sup>)'
    parts = re.split(pattern, md_text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith('<sub>'):
            r = p.add_run(re.sub(r'<sub>|</sub>', '', part)); r.font.subscript = True
        elif part.startswith('<sup>'):
            r = p.add_run(re.sub(r'<sup>|</sup>', '', part)); r.font.superscript = True
        else:
            p.add_run(part)
    return p

def add_image(path, caption, width=6.0):
    try:
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        add_para(f"[IMAGE EMBED FAILED: {path} — {e}]", italic=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10)

def add_table_from_rows(headers, rows, style_name='Light Grid Accent 1'):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = style_name
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(re.sub(r'<[^>]+>|\*\*', '', h)); r.bold = True
        r.font.size = Pt(10)
    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            # Strip markdown bold but show value
            clean = re.sub(r'<sub>([^<]+)</sub>', r'\1', val)
            clean = re.sub(r'<sup>([^<]+)</sup>', r'\1', clean)
            clean = re.sub(r'<br\s*/?>', ' ', clean)
            # Bold if **...** wrapped
            if '**' in clean:
                segments = re.split(r'(\*\*[^*]+\*\*)', clean)
                for seg in segments:
                    if seg.startswith('**') and seg.endswith('**'):
                        r = p.add_run(seg[2:-2]); r.bold = True
                    else:
                        p.add_run(seg)
            else:
                p.add_run(clean)
            for run in p.runs: run.font.size = Pt(10)
    return t

# =================== TITLE PAGE ===================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Mediation of the Educational Gradient in One-Year Acute Exacerbation Risk by Patient-Perceived Symptom Burden in COPD and PRISm: A Prospective Multicentre Korean Cohort Study")
r.bold = True; r.font.size = Pt(14)

doc.add_paragraph()

# Authors
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors = [
    ("Eunjin Kwon", "a,†"),
    ("Won Seo Yoon", "b"),
    ("Ji-Yong Moon", "b"),
    ("Gi Ho Lee", "a"),
    ("Yong-Il Hwang", "c"),
    ("Kwang-Ha Yoo", "b"),
    ("Young-Youl Kim", "a,*"),
    ("Youlim Kim", "b,**"),
]
for i, (name, sup) in enumerate(authors):
    p.add_run(name)
    s = p.add_run(sup); s.font.superscript = True
    if i < len(authors) - 1: p.add_run(", ")

doc.add_paragraph()

# Affiliations
afil = [
    ("a", "Division of Allergy and Respiratory Disease Research, Department of Chronic Disease Convergence Research, National Institute of Health, Cheongju, South Korea"),
    ("b", "Division of Pulmonary and Critical Care Medicine, Department of Internal Medicine, Konkuk University Medical Center, Konkuk University School of Medicine, Seoul, South Korea"),
    ("c", "Division of Pulmonary, Allergy and Critical Care Medicine, Department of Internal Medicine, College of Medicine, Hallym University Sacred Heart Hospital, Anyang, South Korea"),
]
for sup, txt in afil:
    p = doc.add_paragraph()
    s = p.add_run(sup); s.font.superscript = True; s.font.size = Pt(10)
    r = p.add_run(" " + txt); r.font.size = Pt(10)

doc.add_paragraph()

# Footnotes for first author and corresponding
fn = doc.add_paragraph()
s = fn.add_run("† "); s.font.superscript = True; s.font.size = Pt(10)
fn.add_run("First author.").font.size = Pt(10)

fn = doc.add_paragraph()
s = fn.add_run("* "); s.font.superscript = True; s.font.size = Pt(10)
r = fn.add_run("Co-corresponding author: Young-Youl Kim, PhD. Division of Allergy and Respiratory Disease Research, National Institute of Health, 187 Osongsaengmyeong2-ro, Osong-eup, Heungdeok-gu, Cheongju 28159, South Korea. E-mail: youngyk07@korea.kr."); r.font.size = Pt(10)

fn = doc.add_paragraph()
s = fn.add_run("** "); s.font.superscript = True; s.font.size = Pt(10)
r = fn.add_run("Co-corresponding author: Youlim Kim, MD, PhD. Division of Pulmonary and Critical Care Medicine, Konkuk University Medical Center, 120-1 Neungdong-ro, Gwangjin-gu, Seoul 05030, South Korea. E-mail: weilin810707@gmail.com."); r.font.size = Pt(10)

doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run("[Verify before submission] Author list mirrored from the sister paper (Kwon E et al. 2026 Respir Res). Lead author to confirm or amend if KOCOSS investigators differ for this analysis.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x7a, 0x7a, 0x7a)

doc.add_page_break()

# =================== ABSTRACT ===================
add_heading("Abstract", level=1)

abstract_parts = [
    ("Background.", "Lower educational attainment is associated with increased risk of acute exacerbations in chronic obstructive pulmonary disease (COPD), but the intermediate pathways remain poorly characterised within universal-coverage health systems. The mechanistic decomposition of this residual gradient is required to identify upstream policy levers beyond access expansion."),
    ("Methods.", "We analysed 2,712 adults with COPD or preserved ratio impaired spirometry (PRISm) from the Korean COPD Subgroup Study (KOCOSS), a prospective multicentre cohort across 44 hospitals; participants in occupational categories with heterogeneous wage structures (military, housewife, unemployed) were excluded a priori for cohort wage homogeneity. Exposure was three-level educational attainment (Low / Middle / High, reference = High). The outcome was any moderate-to-severe acute exacerbation within one year. The St George's Respiratory Questionnaire (SGRQ) Symptoms domain was designated a priori as the primary candidate mediator; pack-years, SGRQ Activity, and SGRQ Impacts served as prespecified comparator mediators. Total effects were estimated with generalised estimating equations (GEE) logistic regression with hospital clustering, adjusting for age, sex, body mass index, post-bronchodilator FEV1 % predicted, PRISm status, and prior-year exacerbations. Causal mediation followed the VanderWeele 2014 four-way decomposition with 2,000 bootstrap replicates. The hierarchical primary endpoint was the per-tier ordinal trend; the binary Low-versus-High contrast served as a secondary endpoint. A bias-adjusted natural indirect effect (NIE) under a moderate unmeasured-confounding scenario (R=1.10) was reported as a conservative lower bound."),
    ("Results.", "In Model 3, low education was associated with a higher one-year exacerbation risk relative to high education (odds ratio [OR] 1.71, 95% confidence interval [CI] 0.77 to 3.82, p=0.191) with a statistically significant per-tier ordinal trend (OR 1.391, 95% CI 1.05 to 1.84, p-trend=0.020); the extended Y1–Y3 pooled analysis showed OR 1.81 (1.09 to 3.00), p-trend=0.008. In the primary causal mediation analysis (n=850), the SGRQ Symptoms domain was a significant mediator: NIE OR 1.122 (95% CI 1.03 to 1.27), proportion mediated 21.9%. Pack-years and SGRQ Activity did not show significant mediation; SGRQ Impacts showed a borderline NIE (1.06, 95% CI 1.00 to 1.14). Under the conservative R=1.10 adjustment the bias-adjusted NIE was 1.113 (95% CI 1.02 to 1.26). In a GOLD 1–2 sub-cohort (n=605) the NIE was directionally consistent (1.197, 95% CI 1.06 to 1.39); the E-value for the indirect effect was 1.49 (CI lower bound 1.21)."),
    ("Conclusions.", "In a prospective Korean multicentre cohort, approximately one fifth of the educational gradient in one-year COPD or PRISm exacerbation risk was mediated through patient-perceived symptom burden as captured by the SGRQ Symptoms domain, with the indirect effect remaining significant even under conservative confounding adjustment. The finding identifies education as a modifiable upstream determinant operating through symptom experience and supports education-targeted symptom-monitoring care pathways and priority allocation of chronic disease management and pulmonary rehabilitation resources within universal-coverage systems."),
]
for label, text in abstract_parts:
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True
    p.add_run(" " + text)

p = doc.add_paragraph()
r = p.add_run("Keywords: "); r.bold = True
p.add_run("COPD; PRISm; health inequalities; causal mediation analysis; patient-reported outcomes; pulmonary rehabilitation")

doc.add_page_break()

# =================== INTRODUCTION ===================
add_heading("Introduction", level=1)

intro = """Chronic obstructive pulmonary disease (COPD) remains a leading cause of global respiratory morbidity and mortality, with acute exacerbations the principal driver of disease trajectory, healthcare utilisation and population-level burden [1-5]. A consistent observation across high-income systems is that lower educational attainment is associated with increased exacerbation risk, even after access to care is equalised through universal coverage [6-9]. The mechanism through which this residual gradient is transmitted, however, has not been formally decomposed.

Two non-mutually-exclusive families of pathway have been proposed. The first emphasises differential exposure — pack-years, occupational dust, residential pollution — as the channel through which lower socioeconomic position translates into worse lung function and more exacerbations [4, 7]. The second emphasises differential symptom experience and reporting — health literacy, symptom recognition, help-seeking thresholds — as the proximal lever, with patient-reported instruments such as the St George's Respiratory Questionnaire (SGRQ) and the COPD Assessment Test (CAT) capturing the construct [15-18, 27, 28]. Within universal-coverage systems where access has been substantially equalised, the relative contribution of these channels carries direct policy implications: if exposure dominates, expansion of upstream environmental and occupational controls is the lever; if symptom experience dominates, education-targeted symptom-monitoring care pathways become the actionable target.

Causal mediation analysis offers a formal framework for partitioning a total effect into direct and mediated components and for quantifying the proportion mediated under explicit identifying assumptions [20-22]. We applied the VanderWeele four-way decomposition to a prospective Korean multicentre cohort (KOCOSS, 44 hospitals, n=2,712) with the SGRQ Symptoms domain designated a priori as the primary candidate mediator, and pack-years, SGRQ Activity and SGRQ Impacts as prespecified comparator mediators, to test which pathway is the operative channel of the residual education–exacerbation association in the Korean universal-coverage system. We report a hierarchical per-tier ordinal trend as the primary endpoint, a bias-adjusted natural indirect effect under a conservative unmeasured-confounding scenario, and a GOLD 1–2 sub-cohort sensitivity analysis to address potential reverse causation from advanced symptomatic disease.

This work is the second of two pre-planned KOCOSS investigations; the companion paper [37] developed a predictive prognostic index based on the COTE comorbidity score for cardiovascular and exacerbation outcomes. The two analyses share the underlying cohort but address non-overlapping scientific questions on different exposure, outcome and analytic-framework axes."""

for para in intro.split("\n\n"):
    add_md_para(para.strip())

# =================== METHODS ===================
add_heading("Methods", level=1)

add_heading("Study design and cohort", level=2)
add_md_para("This was a secondary analysis of the Korean COPD Subgroup Study (KOCOSS), a prospective multicentre observational cohort of adults with stable COPD or PRISm enrolled across 44 hospitals in Korea. Participants aged ≥40 years with confirmed COPD or PRISm at baseline were eligible. From 2,932 baseline participants, we excluded 220 in occupational categories with heterogeneous wage structures (military KO1_Job=27, n=8; housewife KO1_Job=28, n=124; unemployed KO1_Job=29, n=67; remainder failing other eligibility checks) for cohort wage homogeneity, yielding the canonical analytic cohort of n=2,712. Subjects with missing occupation (KO1_Job=NaN, n=93) were retained in the primary cohort and assessed for sensitivity in S8. Of these, n=850 had complete one-year follow-up plus complete mediator and covariate data and constituted the primary mediation analysis set; the extended Y1–Y3 pooled analysis used n=1,506. The temporal ordering of baseline, mediator and outcome assessment is illustrated in Supplementary Figure S1.")

add_heading("Exposure — educational attainment", level=2)
add_md_para("Educational attainment was self-reported at baseline (KO1_Education) and recoded into three tiers per a priori convention: Low (≤9 years; KO1_Education ∈ {0, 1, 2}), Middle (10–12 years; KO1_Education = 3), and High (>12 years; KO1_Education ∈ {4, 5}). The High tier served as the reference category in regression models. The per-tier ordinal scoring (Low=1, Middle=2, High=3, scaled to a continuous trend variable) supported the primary hierarchical endpoint.")

add_heading("Outcome — one-year acute exacerbation", level=2)
add_md_para("The primary outcome was any moderate-to-severe acute exacerbation within 12 months of baseline, ascertained from prospective KOCOSS follow-up records (FU_Y1_M_exacerbation_YN). Moderate exacerbations required treatment with systemic corticosteroids and/or antibiotics; severe exacerbations required emergency department visit or hospitalisation. The Y1–Y3 extended pooled outcome additionally included exacerbations in months 13–36 from baseline.")

add_heading("Mediators", level=2)
add_md_para("The St George's Respiratory Questionnaire (SGRQ) Symptoms domain was designated a priori as the primary candidate mediator. Pack-years, SGRQ Activity, and SGRQ Impacts served as prespecified comparator mediators. All mediators were measured at baseline.")

add_heading("Covariates", level=2)
add_md_para("Models adjusted for age (continuous), sex (male/female), body mass index (continuous), post-bronchodilator FEV1 % predicted (continuous), PRISm status (binary), and prior-year exacerbations (binary; KO1_P_EX_YN). Inhaled corticosteroid use was excluded as a covariate to avoid treatment-by-indication confounding [11].")

add_heading("Statistical analysis", level=2)
add_md_para("Total effects of education on the one-year exacerbation outcome were estimated using generalised estimating equations (GEE) logistic regression with hospital clustering and an exchangeable working correlation structure [23]. Three nested models were fitted: Model 1 (unadjusted), Model 2 (demographics: age, sex, BMI), and Model 3 (full adjustment, including FEV1 % predicted, PRISm status, and prior-year exacerbations). The hierarchical primary endpoint was the per-tier ordinal trend (continuous education-tier scoring); the binary Low-versus-High contrast served as a secondary endpoint. The Middle-versus-High contrast was reported descriptively. Job=NaN was retained in primary analyses under a missing-at-random (MAR) assumption justified by no differential missingness on outcome or exposure (see Sensitivity §S8); the MCAR assumption was rejected on SGRQ Symptoms (p=0.019).")
add_md_para("Causal mediation analysis followed the VanderWeele 2014 four-way decomposition [20], implemented via the Valeri & VanderWeele 2013 estimator [21] with 2,000 bootstrap replicates for natural direct (NDE) and indirect (NIE) effects and proportion mediated (PM). Sequential ignorability was assumed and is discussed in §Limitations.")
add_md_para("Sensitivity analyses (S1–S8, excluding S6) were prespecified in the protocol; each is reported with its NIE OR, 95% CI and PM (where estimable) in Table 4. Sensitivity S6 (smoker-only stratum, prespecified) was removed post hoc owing to an insufficient stratum size (n=192) that produced a degenerate total-effect bootstrap distribution; this exclusion is reported transparently per ICMJE guidance on prespecified-but-unestimable analyses.")
add_md_para("Quantitative bias assessment. Two bias quantifications were reported. First, the E-value [VanderWeele 2017; Linden 2020] was computed for the indirect effect: 1.49 for the point estimate and 1.21 for the lower confidence bound, indicating the strength of an unmeasured mediator–outcome confounder required to nullify the NIE. Second, a bias-adjusted NIE under a moderate unmeasured-confounding scenario (R=1.10, approximate ρ=0.05) was computed as a conservative lower bound following the mediational E-value framework [Smith & VanderWeele 2019] (Supplementary Table S3), with the threshold value at which the adjusted CI crosses the null (R≈1.21) reported explicitly.")
add_md_para("Code and data availability. Analytic code and the de-identified analysis dataset are deposited at [DOI placeholder, to be minted at acceptance]; the full code repository is available at the corresponding author's institutional repository. All analyses were conducted in Python 3.11 (statsmodels, numpy) and R 4.3 (gee, mediation, CMAverse).")

# Figure 1 (DAG)
add_image(IMG["Figure 1"][0], "Figure 1. Conceptual directed acyclic graph (DAG). Education (exposure) → candidate mediator (SGRQ Symptoms [primary, ★], SGRQ Activity, SGRQ Impacts, or pack-years [comparators]) → one-year acute exacerbation (outcome). Baseline covariates (age, sex, FEV1 % predicted, BMI, PRISm status, prior-year exacerbations) adjust both exposure–mediator and mediator–outcome associations. Time ordering of baseline measurement, mediator assessment and outcome ascertainment is provided in Supplementary Figure S1.", width=5.5)

doc.add_page_break()

# =================== RESULTS ===================
add_heading("Results", level=1)

add_heading("Cohort flow and baseline characteristics", level=2)
add_md_para("Of the 2,932 enrolees in the KOCOSS analytic frame, 220 were excluded under the wage-homogeneity criterion (military 8, housewife 124, unemployed 67; remainder due to age and other eligibility checks), yielding the canonical cohort of n=2,712 (Figure 2). The complete-case mediation analysis set was n=850 (those with complete one-year follow-up plus complete mediator and covariate data); the extended Y1–Y3 pooled analysis used n=1,506. Lower education tiers were older, more often male, had lower FEV1 % predicted, higher prior-year exacerbation prevalence, and higher SGRQ Symptoms scores; smoking exposure (pack-years) also showed a graded pattern (Low 42.4 vs High 33.0, p<0.001), yet did not mediate the education–exacerbation association in formal mediation analysis (see Table 3) (Table 1).")

# Figure 2
add_image(IMG["Figure 2"][0], "Figure 2. Cohort flow diagram. KOCOSS enrolment (n=2,932) → wage-heterogeneous occupations excluded (military, housewife, unemployed) → canonical cohort (n=2,712) → primary mediation analysis set (n=850); extended Y1–Y3 analysis set (n=1,506).", width=5.5)

# Table 1
add_para("Table 1. Baseline characteristics by educational attainment tier, KOCOSS canonical cohort (n=2,712; 59 with missing education excluded from this table; analytic cohort retained).", italic=True, size=10)
t1_headers = ["Variable", "Low (≤9 yrs) n=1,170", "Middle (10–12 yrs) n=1,013", "High (>12 yrs) n=470", "p-value"]
t1_rows = [
    ["Age, years — mean (SD)", "70.5 (7.3)", "67.5 (8.6)", "65.3 (9.9)", "<0.001"],
    ["Male sex, n (%)", "1,048 (89.6)", "940 (92.8)", "451 (96.0)", "<0.001"],
    ["BMI, kg/m² — mean (SD)", "23.4 (3.4)", "23.3 (3.3)", "23.3 (3.4)", "0.96"],
    ["FEV1 % predicted — mean (SD)", "60.1 (18.8)", "62.1 (19.1)", "65.4 (18.5)", "<0.001"],
    ["PRISm, n (%)", "41 (3.5)", "41 (4.0)", "29 (6.2)", "0.049"],
    ["Current smoker, n (%)", "297 (25.4)", "302 (29.8)", "119 (25.4)", "0.043"],
    ["Pack-years — mean (SD)", "42.4 (25.9)", "38.8 (24.2)", "33.0 (19.8)", "<0.001"],
    ["Prior-year exacerbation, n (%)", "199 (17.0)", "124 (12.2)", "52 (11.1)", "0.001"],
    ["SGRQ Symptoms — mean (SD)", "39.6 (21.9)", "34.8 (21.6)", "30.9 (19.9)", "<0.001"],
    ["SGRQ Activity — mean (SD)", "36.5 (26.4)", "32.0 (25.6)", "25.0 (24.3)", "<0.001"],
    ["SGRQ Impacts — mean (SD)", "19.1 (21.0)", "15.4 (19.4)", "12.2 (15.8)", "<0.001"],
    ["Charlson index — median (IQR)", "0.0 (0.0–1.0)", "0.0 (0.0–1.0)", "0.0 (0.0–1.0)", "0.077"],
]
add_table_from_rows(t1_headers, t1_rows)
add_para("Tests: continuous = ANOVA; categorical = chi-square. SGRQ Total row omitted as the three subdomains are reported individually per editorial convention.", italic=True, size=9)

doc.add_paragraph()

add_heading("Total effect of education on one-year exacerbation", level=2)
add_md_para("Per the hierarchical primary endpoint, the per-tier ordinal trend coefficient was statistically significant in Model 3 (per-tier OR 1.391, 95% CI 1.05 to 1.84, p-trend=0.020; Model 2 p-trend=0.014). The secondary binary contrast — low versus high education — yielded OR 1.71 (95% CI 0.77 to 3.82, p=0.191; Table 2). The extended Y1–Y3 pooled outcome analysis showed a stronger and statistically significant Low-versus-High contrast (OR 1.81, 95% CI 1.09 to 3.00, p=0.023; per-tier p-trend=0.008), consistent with a graded exposure–response detectable on the wider follow-up window.")

# Table 2
add_para("Table 2. Adjusted total effect of educational attainment on one-year acute exacerbation risk, KOCOSS canonical cohort.", italic=True, size=10)
t2_headers = ["Comparison", "Model 1 (unadjusted) OR (95% CI)", "Model 2 (demographics) OR (95% CI)", "Model 3 (full adjustment) OR (95% CI)"]
t2_rows = [
    ["Low vs High (secondary)", "1.55 (0.81–2.97)", "1.62 (0.83–3.16)", "1.71 (0.77–3.82), p=0.191"],
    ["Middle vs High (descriptive)", "1.21 (0.61–2.40)", "1.25 (0.62–2.50)", "1.30 (0.59–2.86)"],
    ["Per-tier trend (primary ★)", "1.36 (1.06–1.74)", "1.42 (1.07–1.88), p=0.014", "1.391 (1.05–1.84), p-trend=0.020"],
    ["Y1–Y3 pooled — Low vs High", "—", "—", "1.81 (1.09–3.00), p=0.023; p-trend=0.008"],
]
add_table_from_rows(t2_headers, t2_rows)

doc.add_paragraph()

add_heading("Single-mediator causal mediation analysis", level=2)
add_md_para("In the primary mediation set (n=850), four prespecified mediators were tested in single-mediator models. Only the SGRQ Symptoms domain was a statistically significant mediator (NIE OR 1.122, 95% CI 1.03 to 1.27; PM 21.9%; Table 3). Pack-years and SGRQ Activity did not show significant mediation. SGRQ Impacts contributed a borderline NIE (1.06, 95% CI 1.00 to 1.14) of smaller magnitude than SGRQ Symptoms. The α₁ coefficients (exposure-on-mediator regression for the four candidate mediators) were all positive (+7.897, +5.270, +4.560, +7.596), confirming the expected direction of the education-on-symptom association in fully adjusted models.")

# Table 3
add_para("Table 3. Single-mediator causal mediation analysis of low vs high education on one-year exacerbation risk (n=850).", italic=True, size=10)
t3_headers = ["Mediator", "α₁ (exposure → mediator)", "NIE OR (95% CI)", "NDE OR (95% CI)", "PM (%)", "Significant mediator?"]
t3_rows = [
    ["SGRQ Symptoms (primary ★)", "+7.897", "1.122 (1.03–1.27)", "1.61 (0.74–3.55)", "21.9", "Yes"],
    ["SGRQ Activity", "+5.270", "1.04 (0.97–1.13)", "1.66 (0.75–3.69)", "—", "No"],
    ["SGRQ Impacts", "+4.560", "1.06 (1.00–1.14)", "1.62 (0.73–3.59)", "11.2", "Borderline"],
    ["Pack-years", "+7.596", "1.00 (0.98–1.04)", "1.71 (0.77–3.82)", "—", "No"],
]
add_table_from_rows(t3_headers, t3_rows)

doc.add_paragraph()

add_heading("Sensitivity analyses", level=2)
add_md_para("The primary NIE through SGRQ Symptoms remained significant under the conservative bias-adjusted scenario R=1.10 (adjusted NIE OR 1.113, 95% CI 1.02 to 1.26; Table 4 row 'Bias-adjusted (R=1.10)'; Supplementary Table S3); the threshold value at which the adjusted lower CI crosses 1 was R≈1.21. The corresponding E-value for the indirect effect was 1.49 (lower CI bound 1.21). Direction was preserved in the GOLD 1–2 mild-to-moderate sub-cohort (n=605; NIE 1.197, 95% CI 1.06 to 1.39; PM 43.1%; Supplementary Table S4), in which baseline measurement is least susceptible to reverse-causation bias from advanced symptomatic disease. The Y1–Y3 extended pooled analysis showed a stronger total effect (OR 1.81, 95% CI 1.09 to 3.00). Excluding the n=93 occupation-missing participants (S8) yielded an OR of 1.77 (0.78 to 4.02) and a per-tier trend of 1.39 (1.04 to 1.85, p=0.026), consistent with the primary analysis under MAR. Inverse-probability weighting for missingness on the mediation set (S7, n=850) yielded NIE OR 1.149 (1.05 to 1.29), marginally stronger than the primary; additional Charlson adjustment (S2, n=724) yielded a borderline NIE of 1.100 (1.00 to 1.24).")

# Table 4
add_para("Table 4. Sensitivity analyses of the SGRQ Symptoms–mediated effect of low vs high education on one-year exacerbation.", italic=True, size=10)
t4_headers = ["Analysis", "n", "NIE OR (95% CI)", "TE / Low–High OR (95% CI)", "PM (%)"]
t4_rows = [
    ["Primary", "850 (med set)", "1.122 (1.03–1.27) ★", "1.71 (0.77–3.82)", "21.9"],
    ["Bias-adjusted (R=1.10)", "850", "1.113 (1.02–1.26) ★", "—", "—"],
    ["S1 FEV1 unadjusted", "1,415", "—", "1.82 (0.86–3.85)", "—"],
    ["S5 COPD-only", "1,306", "—", "1.60 (0.78–3.27)", "—"],
    ["S2 + Charlson", "724", "1.100 (1.00–1.24) (borderline)", "1.668 (0.94–3.60)", "18.6"],
    ["S7 IPW for missingness", "850", "1.149 (1.05–1.29) ★", "1.878 (1.09–4.14)", "22.0"],
    ["Y1–Y3 pooled", "1,506", "—", "1.81 (1.09–3.00), p-trend=0.008", "—"],
    ["S8 Job=NaN excluded", "1,328", "—", "1.77 (0.78–4.02); per-tier OR 1.39 (1.04–1.85), p=0.026", "—"],
    ["GOLD 1–2 sub-cohort (Supp S4)", "605", "1.197 (1.06–1.39) ★", "1.531 (0.755–3.105)", "43.1"],
]
add_table_from_rows(t4_headers, t4_rows)
add_para("★ = 95% CI excludes 1.", italic=True, size=9)

# Figure 3
add_image(IMG["Figure 3"][0], "Figure 3. Forest plot of the natural indirect effect (NIE) of education through the SGRQ Symptoms domain across primary, bias-adjusted, and seven sensitivity analyses (n=2,712 cohort). Horizontal bars show 95% confidence intervals; the vertical dashed line marks the null (OR = 1). The primary estimate is annotated ★. S6 (current-smokers-only stratum) was removed post hoc owing to insufficient stratum size (n=192) producing a degenerate total-effect bootstrap; transparent reporting per ICMJE.", width=6.0)

doc.add_page_break()

# =================== DISCUSSION ===================
add_heading("Discussion", level=1)

add_heading("Principal findings", level=2)
add_md_para("In this prospective multicentre Korean cohort of individuals with COPD or PRISm, low educational attainment was associated with a graded, statistically significant per-tier increase in one-year acute exacerbation risk (per-tier OR 1.391, p-trend=0.020), with the effect strengthening on the extended Y1–Y3 outcome window (OR 1.81, p-trend=0.008). Formal causal mediation identified the SGRQ Symptoms domain as the primary mediating pathway: NIE OR 1.122 (95% CI 1.03 to 1.27), accounting for approximately one fifth of the total effect. The indirect effect remained significant under conservative confounding adjustment (R=1.10; adjusted NIE 1.113, 95% CI 1.02 to 1.26) and was directionally preserved in a GOLD 1–2 mild-to-moderate sub-cohort (NIE 1.197, 95% CI 1.06 to 1.39), in which baseline assessment is least susceptible to reverse-causation bias from advanced disease. Pack-years and SGRQ Activity did not show significant mediation; SGRQ Impacts contributed a smaller borderline indirect effect. This pattern suggests that within a universal-coverage system, symptom experience — rather than exposure intensity or activity limitation — is the operative pathway linking education to exacerbation risk.")

add_heading("Comparison with prior literature and direction triangulation", level=2)
add_md_para("Our directional finding — that lower-education COPD populations carry higher exacerbation risk — is consistent with cohort-based observations from the Danish Glostrup [7], Canadian ICES [6, 8], and UK CPRD [Sin et al., as discussed below] datasets. The novel contribution is the formal mediation decomposition through patient-perceived symptom burden. The directional consistency of SES → COPD outcome associations across published cohorts with non-overlapping unmeasured-confounder profiles — UK Biobank [30] and the ECLIPSE cohort, where pack-years, comorbidity and lung-function adjustments use distinct measurement instruments — provides external triangulation that strengthens the causal interpretation beyond what single-cohort sensitivity analysis can establish.")

add_heading("Possible mechanisms", level=2)
add_md_para("Three non-mutually-exclusive mechanisms may underlie the SGRQ Symptoms mediation pathway. First, lower educational attainment is associated with reduced health literacy [27, 28], which may delay symptom recognition and help-seeking. Second, lower-education populations may experience differential symptom amplification driven by comorbidity, medication adherence variability, and care-pathway navigation difficulties, which the SGRQ Symptoms domain captures more sensitively than activity-limitation or impact-on-life domains [15-18]. Third, the symptom-reporting threshold itself may be calibrated differently across educational tiers, with lower-tier respondents experiencing or reporting symptoms at lower clinical severity and translating into earlier exacerbation events.")

add_heading("Strengths", level=2)
add_md_para("Key strengths include a large prospective multicentre cohort (n=2,712, 44 sites) with an a priori wage-homogeneity exclusion, standardised ascertainment of exposures, mediators, covariates and one-year exacerbations, and a formal four-way causal mediation framework with 2,000 bootstrap replicates. The robustness strategy spans bias-adjusted lower bounds, an E-value, a GOLD 1–2 sub-cohort sensitivity addressing reverse causation, an MAR analysis for occupation missingness, and an extended Y1–Y3 pooled outcome window. The S7 inverse-probability-weighted analysis (n=850) yielded NIE OR 1.149 (1.05 to 1.29), marginally stronger than the primary estimate (1.122) — strengthening confidence that the indirect effect is not an artefact of complete-case selection. The S2 additional Charlson adjustment produced a borderline NIE 1.100 (1.00 to 1.24); the lower CI bound at 1.00 indicates that residual comorbidity confounding could partially explain the SGRQ-mediated effect, but direction was preserved.")

add_heading("Limitations", level=2)
add_md_para("Several limitations should be considered. First, the E-value of 1.49 indicates that the indirect effect is sensitive to moderate unmeasured confounding (R≈1.21 nullification threshold); plausible KOCOSS unmeasured factors include occupational dust history, residential PM2.5 exposure and family COPD severity, any of which could approach this threshold, and the bias-adjusted NIE (R=1.10) is therefore reported as a conservative lower bound. Second, education and occupation were ascertained at baseline post-COPD-diagnosis; reverse causation from symptomatic disease cannot be fully excluded despite the GOLD 1–2 sub-cohort robustness. Third, occupation was missing for 93 participants (Job=NaN) and the MCAR assumption was rejected on SGRQ Symptoms (p=0.019), although the outcome and exposure showed no missingness-related differential and the S8 sensitivity analysis was consistent with the primary; analyses are therefore interpreted under a MAR assumption. Fourth, the causal mediation framework relies on a sequential ignorability assumption that cannot be verified empirically, and our findings are presented as one of multiple plausible interpretations rather than as definitive proof of causation. Fifth, although the SGRQ Symptoms domain and our exacerbation outcome draw on overlapping symptom constructs, complete construct equivalence is unlikely given the temporal, qualitative and operational distinctions between chronic symptom burden and an acute episode requiring therapeutic escalation.")

add_heading("Implications for policy and population health", level=2)
add_md_para("Within a universal-coverage system, the residual one-year exacerbation gradient by educational attainment — approximately one fifth of which is channelled through patient-perceived symptom burden — points away from access expansion as the principal lever and identifies education itself as a modifiable upstream determinant operating through symptom experience. Education-targeted public-health and care-system interventions follow as the actionable upstream pathway: priority allocation of chronic disease management programmes and pulmonary rehabilitation to lower-education COPD strata, integration of patient-reported symptom instruments such as the SGRQ or the COPD Assessment Test into the routine outpatient encounter for these strata, and the development of symptom-monitoring care pathways that lower the symptom-recognition and help-seeking threshold for lower-education populations. The downstream feasibility of such differential allocation has been strengthened by emerging evidence on telerehabilitation and digital self-management programmes [Cox 2022; Gloeckl 2025; Bourne 2017], which lower the resource intensity of population-level deployment. The absence of pack-years mediation in our data does not imply that conventional tobacco-control measures should be displaced; rather, education-targeted symptom-monitoring care should be added alongside rather than substituted for them. These implications are hypothesis-generating and require confirmation in implementation studies; nonetheless, addressing how lower-education COPD populations experience and report symptoms — rather than merely whether they reach care — may be the higher-yield population lever in universal-coverage systems.")

add_heading("Conclusion", level=2)
add_md_para("In a prospective Korean multicentre cohort, a graded educational gradient in one-year COPD or PRISm exacerbation risk was identified, approximately one fifth of which was mediated through patient-perceived symptom burden as captured by the SGRQ Symptoms domain. The indirect effect remained significant under conservative unmeasured-confounding adjustment and was directionally preserved in a mild-to-moderate sub-cohort. These findings identify education as a modifiable upstream determinant operating through symptom experience and support education-targeted symptom-monitoring care pathways and priority resource allocation as upstream policy levers within universal-coverage COPD systems.")

doc.add_page_break()

# =================== SUPPLEMENTARY MATERIAL ===================
add_heading("Supplementary material", level=1)
add_md_para("Supplementary Figure S1 — temporal ordering of baseline measurement, mediator assessment and outcome ascertainment in KOCOSS (below). Supplementary Table S3 — quantitative bias assessment (E-value sweep, ρ ∈ {0.05, 0.10, 0.15, 0.20, 0.25}); Supplementary Table S4 — GOLD 1–2 sub-cohort sensitivity (n=605); Supplementary Table S8 — Job-NaN MAR test (Little's test plus per-characteristic comparison). These tables are available as separate files (S3_ignorability_sensitivity, S4_gold12_subcohort, S8_mcar_test) in machine-readable JSON and Markdown formats.")

add_image(IMG["Supp S1"][0], "Supplementary Figure S1. Temporal ordering of baseline measurement (T0), mediator assessment, and outcome ascertainment (T+12 months for primary; T+12 to T+36 months for extended Y1–Y3 pooled analysis) in the KOCOSS cohort. Education and occupation were ascertained at baseline; reverse causation from advanced symptomatic disease cannot be fully excluded but is addressed in the GOLD 1–2 sub-cohort sensitivity analysis (Table 4 / Supplementary Table S4).", width=6.0)

doc.add_page_break()

# =================== DECLARATIONS ===================
add_heading("Declarations", level=1)

decls = [
    ("Funding.", "This work was supported by the Research Program funded by Korea National Institute of Health (grant numbers 2016ER670100, 2016ER670101, 2016ER670102, 2018ER670100, 2018ER670101, 2018ER670102, 2021ER120500, 2021ER120501, 2021ER120502, 2024ER120100, 2024ER120101, 2024ER120102, and 2023NI00702)."),
    ("Competing interests.", "The authors declare no competing interests."),
    ("Ethics approval and consent to participate.", "All hospitals participating in the Korean COPD Subgroup Study (KOCOSS) obtained approval from the relevant Institutional Review Boards, including Konkuk University Medical Center (IRB No. KHH1010338). All participants provided written informed consent."),
    ("Consent for publication.", "Not applicable."),
    ("Data availability.", "Analytic code and the de-identified analysis dataset will be deposited at Zenodo with a citable DOI prior to peer review (DOI placeholder — to be inserted at submission; see Methods §Statistical analysis). Raw KOCOSS data are subject to the KOCOSS data-sharing policy and relevant Korean regulations; requests will be reviewed by the corresponding authors."),
    ("Acknowledgements.", "We thank the participants and investigators of the Korean COPD Subgroup Study (KOCOSS)."),
    ("Author contributions.", "Eunjin Kwon: Conceptualization, Methodology, Data curation, Formal analysis, Visualization, Validation, Writing — original draft, Writing — review & editing. Won Seo Yoon: Methodology, Data curation, Validation, Writing — review & editing. Ji-Yong Moon: Investigation, Writing — review & editing. Gi Ho Lee: Data curation, Writing — review & editing. Yong-Il Hwang: Investigation, Data curation, Writing — review & editing. Kwang-Ha Yoo: Investigation, Supervision, Writing — review & editing. Young-Youl Kim: Conceptualization, Supervision, Funding acquisition, Project administration, Writing — review & editing. Youlim Kim: Conceptualization, Supervision, Project administration, Writing — review & editing."),
]
for label, text in decls:
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True
    p.add_run(" " + text)

note = doc.add_paragraph()
r = note.add_run("[Verify before submission] Funding/Ethics/Competing interests mirrored from sister paper (Kwon E et al. 2026 Respir Res). Author contributions follow a standard CRediT pattern observed in the sister paper for overlapping authors. Lead author to confirm or amend per actual contribution.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x7a, 0x7a, 0x7a)

doc.add_page_break()

# =================== REFERENCES ===================
add_heading("References", level=1)

refs = [
    "MacLeod M, Papi A, Contoli M, et al. Chronic obstructive pulmonary disease exacerbation fundamentals: diagnosis, treatment, prevention and disease impact. Respirology 2021;26:532–551. doi:10.1111/resp.14041",
    "Global Initiative for Chronic Obstructive Lung Disease. Global Strategy for Prevention, Diagnosis and Management of COPD: 2025 Report. [verify — goldcopd.org]",
    "GBD 2019 Chronic Respiratory Diseases Collaborators. Prevalence and attributable health burden of chronic respiratory diseases, 1990–2017. Lancet Respir Med 2020;8:585–596. doi:10.1016/S2213-2600(20)30105-3",
    "Yang IA, Jenkins CR, Salvi SS. Chronic obstructive pulmonary disease in never-smokers: risk factors, pathogenesis, and implications. Lancet Respir Med 2022;10:497–511. doi:10.1016/S2213-2600(21)00506-3",
    "Stolz D, Mkorombindo T, Schumann DM, et al. Towards the elimination of chronic obstructive pulmonary disease: a Lancet Commission. Lancet 2022;400:921–972. doi:10.1016/S0140-6736(22)01273-9",
    "Gershon AS, Dolmage TE, Stephenson A, Jackson B. COPD and socioeconomic status: a systematic review. COPD 2012;9:216–226. doi:10.3109/15412555.2011.648030 [verify]",
    "Prescott E, Lange P, Vestbo J. Socioeconomic status, lung function and admission to hospital for COPD. Eur Respir J 1999;13:1109–1114. [verify]",
    "Marmot M. Social determinants of health inequalities. Lancet 2005;365:1099–1104. doi:10.1016/S0140-6736(05)71146-6",
    "Galobardes B, Shaw M, Lawlor DA, Lynch JW, Davey Smith G. Indicators of socioeconomic position (part 1). J Epidemiol Community Health 2006;60:7–12. doi:10.1136/jech.2004.023531",
    "Wan ES, Fortis S, Regan EA, et al. Epidemiology, genetics, and subtyping of preserved ratio impaired spirometry (PRISm) in COPDGene. Respir Res 2014;15:89. [verify]",
    "Bhatt SP, Abadi E, Anzueto A, et al. A multidimensional diagnostic approach for COPD. JAMA 2025;333:2164–2175. doi:10.1001/jama.2025.7358",
    "Hurst JR, Vestbo J, Anzueto A, et al. Susceptibility to exacerbation in COPD. N Engl J Med 2010;363:1128–1138. doi:10.1056/NEJMoa0909883",
    "Celli BR, Fabbri LM, Aaron SD, et al. An updated definition and severity classification of COPD exacerbations: the Rome Proposal. Am J Respir Crit Care Med 2021;204:1251–1258. doi:10.1164/rccm.202108-1819PP",
    "Divo M, Cote C, de Torres JP, et al. Comorbidities and risk of mortality in COPD. Am J Respir Crit Care Med 2012;186:155–161. doi:10.1164/rccm.201201-0034OC",
    "Jones PW, Quirk FH, Baveystock CM, Littlejohns P. A self-complete measure of health status for chronic airflow limitation: the SGRQ. Am Rev Respir Dis 1992;145:1321–1327. doi:10.1164/ajrccm/145.6.1321",
    "Jones PW. SGRQ: MCID. COPD 2005;2:75–79. doi:10.1081/COPD-200050513",
    "Kessler R, Partridge MR, Miravitlles M, et al. Symptom variability in patients with severe COPD. Eur Respir J 2011;37:264–272. doi:10.1183/09031936.00051110",
    "Choi JY, Yoon HK, Shin KC, et al. CAT score and SGRQ definitions of chronic bronchitis. Int J Chron Obstruct Pulmon Dis 2019;14:3043–3052. doi:10.2147/COPD.S228307",
    "Jang JG, Kim Y, Lee JK, et al. Clinical characteristics of individuals with COPD, pre-COPD and smokers with normal lung function in Korea. Tuberc Respir Dis (Seoul) 2025;89:75–85. doi:10.4046/trd.2025.0040",
    "VanderWeele TJ. A unification of mediation and interaction: a 4-way decomposition. Epidemiology 2014;25:749–761. doi:10.1097/EDE.0000000000000121",
    "Valeri L, VanderWeele TJ. Mediation analysis allowing for exposure–mediator interactions and causal interpretation. Psychol Methods 2013;18:137–150. doi:10.1037/a0031034",
    "Imai K, Keele L, Tingley D. A general approach to causal mediation analysis. Psychol Methods 2010;15:309–334. doi:10.1037/a0020761",
    "Zeger SL, Liang KY. Longitudinal data analysis for discrete and continuous outcomes. Biometrics 1986;42:121–130.",
    "Seaman SR, White IR. Review of inverse probability weighting for dealing with missing data. Stat Methods Med Res 2013;22:278–295. doi:10.1177/0962280210395740",
    "VanderWeele TJ, Ding P. Sensitivity analysis in observational research: introducing the E-value. Ann Intern Med 2017;167:268–274. doi:10.7326/M16-2607",
    "Linden A, Mathur MB, VanderWeele TJ. Conducting sensitivity analysis for unmeasured confounding using E-values. Stata J 2020;20:162–175. [verify]",
    "Nutbeam D. The evolving concept of health literacy. Soc Sci Med 2008;67:2072–2078. doi:10.1016/j.socscimed.2008.09.050",
    "Berkman ND, Sheridan SL, Donahue KE, Halpern DJ, Crotty K. Low health literacy and health outcomes: an updated systematic review. Ann Intern Med 2011;155:97–107. doi:10.7326/0003-4819-155-2-201107190-00005",
    "Effing TW, Vercoulen JH, Bourbeau J, et al. Definition of a COPD self-management intervention: International Expert Group consensus. Eur Respir J 2016;48:46–54. doi:10.1183/13993003.00025-2016",
    "Han YY, Yan Q, Chen W, Celedón JC. Child maltreatment, anxiety and depression, and asthma among British adults in the UK Biobank. Eur Respir J 2022;60(4):2103160. doi:10.1183/13993003.03160-2021",
    "Jones PW, Rutten-van Mölken MPMH, Agusti A, et al. Reporting patient-reported outcomes in COPD: a narrative review. Eur Respir J 2019;54:1900168. [verify]",
    "Lamberton CE, Mosher CL. Review of the evidence for pulmonary rehabilitation in COPD. Respir Care 2024;69:686–696. doi:10.4187/respcare.11541",
    "Spruit MA, Singh SJ, Garvey C, et al. ATS/ERS statement: key concepts and advances in pulmonary rehabilitation. Am J Respir Crit Care Med 2013;188:e13–e64. doi:10.1164/rccm.201309-1634ST",
    "Cox NS, McDonald CF, Mahal A, et al. Telerehabilitation for chronic respiratory disease: a randomised controlled equivalence trial. Thorax 2022;77:643–651. doi:10.1136/thoraxjnl-2021-216934",
    "Gloeckl R, Spielmanns M, Stankeviciene A, et al. Smartphone application-based pulmonary rehabilitation in COPD. Thorax 2025;80:209–217. doi:10.1136/thorax-2024-221803",
    "Bourne S, DeVos R, North M, et al. Online versus face-to-face pulmonary rehabilitation for patients with COPD. BMJ Open 2017;7:e014580. doi:10.1136/bmjopen-2016-014580",
    "Kwon E, Yoon WS, Moon JY, Lee GH, Hwang YI, Yoo KH, Kim YY, Kim Y. Cardiovascular components of the COTE index predict acute exacerbations and healthcare costs in patients with chronic obstructive pulmonary disease: a nationwide linked cohort study. Respir Res 2026;27. doi:10.1186/s12931-026-03677-4",
    "Smith LH, VanderWeele TJ. Mediational E-values: approximate sensitivity analysis for unmeasured mediator–outcome confounding. Epidemiology 2019;30:835–837. doi:10.1097/EDE.0000000000001064",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. "); r.bold = True
    p.add_run(ref).font.size = Pt(10)

# Save
doc.save(OUT)
print(f"DONE: {OUT}")
print(f"Size: {OUT.stat().st_size} bytes")
print(f"Refs count: {len(refs)}")

