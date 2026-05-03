"""
Build Korean Word .docx from ver3.5_F-rev2 manuscript.
한글 번역 + 동일 표·그림 embed.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("C:/Projects/legend-team/reports/2026-05-02_copd-paper-part2-residual")
OUT = ROOT / "Manuscript_ver3.5_F-rev3_K.docx"

FIGS_V2 = "C:/Projects/legend-team/reports/2026-05-03_copd-figure1-consort-redesign/figures_v2"
IMG = {
    "Figure 1": "C:/Projects/COPD/Fig. 1 (2).jpg",
    "Figure 2": f"{FIGS_V2}/Fig1_CONSORT_v3.png",
    "Figure 3": f"{FIGS_V2}/Fig3_forest_NIE_v2.png",
    "Supp S1": "C:/Projects/COPD/Data/ver2.0/submission_session174_2026-05-02/02_analysis_outputs/S1_time_ordering.png",
    "Supp S2": f"{FIGS_V2}/Fig2_followup_sweep_v2.png",
}

doc = Document()
sec = doc.sections[0]
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5);  sec.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(11)

def set_kr_font(run, name='Malgun Gothic'):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1f, 0x40, 0x68)
        set_kr_font(r)
    return h

def add_para(text, italic=False, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    if italic: r.italic = True
    if bold: r.bold = True
    if size: r.font.size = Pt(size)
    set_kr_font(r)
    return p

def add_para_bold_label(label, body, size=11):
    p = doc.add_paragraph()
    r1 = p.add_run(label); r1.bold = True; r1.font.size = Pt(size); set_kr_font(r1)
    r2 = p.add_run(" " + body); r2.font.size = Pt(size); set_kr_font(r2)
    return p

def add_image(path, caption, width=6.0):
    try:
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        add_para(f"[이미지 임베드 실패: {path} — {e}]", italic=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10); set_kr_font(r)

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10); set_kr_font(r)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            if '**' in val:
                segments = re.split(r'(\*\*[^*]+\*\*)', val)
                for seg in segments:
                    if seg.startswith('**') and seg.endswith('**'):
                        r = p.add_run(seg[2:-2]); r.bold = True; r.font.size = Pt(10); set_kr_font(r)
                    else:
                        r = p.add_run(seg); r.font.size = Pt(10); set_kr_font(r)
            else:
                r = p.add_run(val); r.font.size = Pt(10); set_kr_font(r)
    return t

# =================== TITLE PAGE ===================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("교육수준에 따른 1년 급성악화 위험의 환자 인지 증상 부담을 통한 매개효과: COPD 및 PRISm 환자 대상 한국 다기관 전향적 코호트 연구")
r.bold = True; r.font.size = Pt(14); set_kr_font(r)

doc.add_paragraph()
en_subtitle = doc.add_paragraph()
en_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = en_subtitle.add_run("(Mediation of the Educational Gradient in One-Year Acute Exacerbation Risk by Patient-Perceived Symptom Burden in COPD and PRISm: A Prospective Multicentre Korean Cohort Study)")
r.italic = True; r.font.size = Pt(11); set_kr_font(r)

doc.add_paragraph()

# Authors (영문 표기 유지 — 학술논문 관행)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors = [
    ("Eunjin Kwon (권은진)", "a,†"),
    ("Won Seo Yoon", "b"),
    ("Ji-Yong Moon", "b"),
    ("Gi Ho Lee", "a"),
    ("Yong-Il Hwang", "c"),
    ("Kwang-Ha Yoo", "b"),
    ("Young-Youl Kim", "a,*"),
    ("Youlim Kim", "b,**"),
]
for i, (name, sup) in enumerate(authors):
    r = p.add_run(name); set_kr_font(r)
    s = p.add_run(sup); s.font.superscript = True; set_kr_font(s)
    if i < len(authors) - 1: p.add_run(", ")

doc.add_paragraph()

afil = [
    ("a", "질병관리청 국립보건연구원 만성질환융합연구부 알레르기·호흡기질환연구과 (Division of Allergy and Respiratory Disease Research, Department of Chronic Disease Convergence Research, National Institute of Health, Cheongju, South Korea)"),
    ("b", "건국대학교 의과대학 건국대학교병원 호흡기내과 (Division of Pulmonary and Critical Care Medicine, Konkuk University Medical Center, Seoul, South Korea)"),
    ("c", "한림대학교 의과대학 한림대학교성심병원 호흡기·알레르기내과 (Division of Pulmonary, Allergy and Critical Care Medicine, Hallym University Sacred Heart Hospital, Anyang, South Korea)"),
]
for sup, txt in afil:
    p = doc.add_paragraph()
    s = p.add_run(sup); s.font.superscript = True; s.font.size = Pt(10); set_kr_font(s)
    r = p.add_run(" " + txt); r.font.size = Pt(10); set_kr_font(r)

doc.add_paragraph()
add_para("† 제1저자.", size=10)
add_para("* 공동교신저자: Young-Youl Kim, PhD. 187 오송생명2로, 오송읍, 흥덕구, 청주 28159, 대한민국. E-mail: youngyk07@korea.kr.", size=10)
add_para("** 공동교신저자: Youlim Kim, MD, PhD. 120-1 능동로, 광진구, 서울 05030, 대한민국. E-mail: weilin810707@gmail.com.", size=10)

doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run("[제출 전 확인 필요] 저자 명단은 자매 논문(Kwon E et al. 2026 Respir Res)을 기준으로 작성됨. KOCOSS 연구진 구성이 본 분석에서 다를 경우 제1저자가 확정·수정 필요.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x7a, 0x7a, 0x7a); set_kr_font(r)

doc.add_page_break()

# =================== ABSTRACT ===================
add_heading("초록 (Abstract)", level=1)

add_para_bold_label("배경 (Background).", "낮은 교육수준은 만성폐쇄성폐질환(COPD)에서 급성악화 위험 증가와 연관되지만, 보편적 의료보장 체계 내에서 그 중간 경로는 잘 규명되지 않았다. 잔여 gradient의 기전적 분해는 의료 접근성 확대 외의 상위 정책 lever 식별을 위해 필요하다.")

add_para_bold_label("방법 (Methods).", "한국 COPD 코호트 연구(KOCOSS, 44개 병원)에 등록된 COPD 또는 보존비율 폐기능저하(PRISm) 환자 2,712명을 분석하였다. 임금 구조가 이질적인 직업군(군인·주부·무직)은 코호트 임금 동질성을 위해 사전 제외하였다. 노출 변수는 3단계 교육수준(저[≤9년]/중[10–12년]/고[>12년], 기준=고)이었다. 결과 변수는 1년 이내 중등도 이상 급성악화 발생이었다. St George's 호흡기 설문(SGRQ) 증상 영역을 사전에 일차 매개변수로 지정하였고, 흡연력(pack-years)·SGRQ 활동·SGRQ 영향 영역을 사전 비교 매개변수로 설정하였다. 총효과는 병원을 군집변수로 한 일반화 추정방정식(GEE) 로지스틱 회귀로 추정하였으며 연령·성별·체질량지수·BD 후 FEV1 % 예측치·PRISm 상태·전년도 악화 이력을 보정하였다. 인과 매개분석은 VanderWeele 2014 4-way decomposition을 따랐으며 부트스트랩 2,000회를 적용하였다. 일차 위계 종점은 단계별 순서형 추세였고, 저-vs-고 이항 비교는 이차 종점이었다. 중간 정도의 미측정 교란(R=1.10) 시나리오 하 편향보정 자연간접효과(NIE)를 보수적 하한으로 보고하였다.")

add_para_bold_label("결과 (Results).", "Model 3에서 저학력군은 고학력군 대비 1년 악화 위험이 높았으나(OR 1.71, 95% CI 0.77~3.82, p=0.191) 단계별 순서형 추세는 통계적으로 유의하였다(OR 1.391, 95% CI 1.05~1.84, p-trend=0.020). Y1–Y3 확장 분석에서는 OR 1.81 (1.09~3.00), p-trend=0.008이었다. 일차 인과 매개분석(n=850)에서 SGRQ 증상 영역은 유의한 매개변수였다: NIE OR 1.122 (95% CI 1.03~1.27), 매개비율(PM) 21.9%. Pack-years와 SGRQ 활동은 유의한 매개효과를 보이지 않았으며, SGRQ 영향은 경계적 NIE를 보였다(1.06, 95% CI 1.00~1.14). 보수적 R=1.10 보정 하 편향보정 NIE는 1.113 (95% CI 1.02~1.26)이었다. GOLD 1–2 하위 코호트(n=605)에서 NIE는 방향적으로 일치하였다(1.197, 95% CI 1.06~1.39). 간접효과의 E-값은 1.49(CI 하한 1.21)였다.")

add_para_bold_label("결론 (Conclusions).", "한국 다기관 전향적 코호트에서 1년 COPD 또는 PRISm 악화 위험의 교육수준 gradient 중 약 1/5이 SGRQ 증상 영역으로 측정된 환자 인지 증상 부담을 통해 매개되었으며, 보수적 교란보정 하에서도 간접효과는 유의하였다. 본 연구는 교육수준을 증상 경험을 통해 작용하는 수정 가능한 상위 결정요인으로 식별하며, 보편적 의료보장 체계 내 교육 수준별 증상 모니터링 진료 경로 및 만성질환 관리·폐재활 자원의 우선 배정을 정책적 lever로 지지한다.")

p = doc.add_paragraph()
r = p.add_run("Keywords: "); r.bold = True; set_kr_font(r)
r = p.add_run("COPD; PRISm; 건강 불평등; 인과 매개분석; 환자보고 결과지표; 폐재활"); set_kr_font(r)

doc.add_page_break()

# =================== INTRO ===================
add_heading("서론 (Introduction)", level=1)

intros = [
    "만성폐쇄성폐질환(COPD)은 전 세계 호흡기 질환 이환·사망의 주요 원인으로 지속되고 있으며, 급성악화는 질병 경과·의료이용·인구집단 부담의 주된 동인이다 [1-5]. 고소득 의료체계에서 일관되게 관찰되는 현상은, 보편적 의료보장으로 접근성이 평준화된 이후에도 낮은 교육수준이 악화 위험 증가와 연관되는 것이다 [6-9]. 그러나 이 잔여 gradient가 어떤 기전을 통해 전달되는지는 공식적으로 분해된 적이 없다.",
    "두 가지 상호 배타적이지 않은 경로 군이 제안되어 왔다. 첫째는 차등 노출(흡연량·직업적 분진·거주지 대기오염)이 낮은 사회경제적 위치를 통해 폐기능 저하 및 악화로 전이되는 채널을 강조한다 [4, 7]. 둘째는 차등 증상 경험·보고(건강 문해력·증상 인지·도움요청 임계)를 근위 lever로 강조하며, SGRQ나 COPD Assessment Test(CAT) 같은 환자보고 도구가 그 구성개념을 포착한다 [15-18, 27, 28]. 접근성이 상당히 평준화된 보편적 의료보장 체계 내에서 이 두 채널의 상대적 기여도는 직접적인 정책적 함의를 가진다: 노출이 우세하면 환경·직업 통제 확대가 lever이며, 증상 경험이 우세하면 교육 수준별 증상 모니터링 진료 경로가 실행 가능 표적이 된다.",
    "인과 매개분석은 명시적 식별 가정 하에 총효과를 직접·매개 성분으로 분할하고 매개비율을 정량화하는 공식적 framework을 제공한다 [20-22]. 본 연구는 한국 다기관 전향적 코호트(KOCOSS, 44개 병원, n=2,712)에 VanderWeele 4-way decomposition을 적용하였다. SGRQ 증상 영역을 사전에 일차 매개변수로, pack-years·SGRQ 활동·SGRQ 영향 영역을 사전 비교 매개변수로 설정하여, 한국 보편적 의료보장 체계의 잔여 교육-악화 연관에서 어느 경로가 작동 채널인지 검정하였다. 일차 종점은 위계적 단계별 순서형 추세를, 보수적 미측정 교란 시나리오 하 편향보정 자연간접효과를, 그리고 진행성 증상 질환에서의 잠재적 역인과 우려를 다루기 위한 GOLD 1–2 하위 코호트 민감도 분석을 보고한다.",
    "본 연구는 KOCOSS 사전 계획된 두 분석 중 두 번째이다. 자매 논문 [37]은 복합 동반질환 지수 기반 심혈관 및 악화 결과 예측 prognostic index를 개발하였다. 두 분석은 동일 코호트를 공유하지만 노출·결과·분석 framework 세 축에서 비중복 과학적 질문을 다룬다.",
]
for txt in intros:
    add_para(txt)

# =================== METHODS ===================
add_heading("방법 (Methods)", level=1)

add_heading("연구 설계 및 코호트", level=2)
add_para("본 연구는 한국 COPD 코호트 연구(KOCOSS) — 한국 44개 병원에서 등록된 안정형 COPD 또는 PRISm 성인 환자 대상 다기관 전향적 관찰 코호트 — 의 이차 분석이다. 베이스라인에서 COPD 또는 PRISm이 확인된 40세 이상 환자가 적격이었다. 베이스라인 2,932명 중 임금 구조가 이질적인 직업 범주(군인 KO1_Job=27, n=8; 주부 KO1_Job=28, n=124; 무직 KO1_Job=29, n=67; 그 외 적격 기준 미충족)에 속한 220명을 코호트 임금 동질성을 위해 제외하여, 최종 분석 코호트 n=2,712를 구성하였다. 직업 결측(KO1_Job=NaN, n=93)인 대상자는 일차 코호트에 유지되고 S8 민감도에서 평가되었다. 이 중 1년 추적 완료 + 매개변수·공변수 완전 데이터를 가진 850명이 일차 매개분석 세트가 되었으며, Y1–Y3 확장 풀링 분석은 n=1,506을 사용하였다. 베이스라인·매개변수·결과 측정의 시간 순서는 보충 그림 S1에 도시하였다.")

add_heading("노출 변수 — 교육수준", level=2)
add_para("교육수준은 베이스라인에서 자가 보고(KO1_Education)되었으며 사전 정의에 따라 3단계로 재코딩하였다: 저(≤9년; KO1_Education ∈ {0, 1, 2}), 중(10–12년; KO1_Education = 3), 고(>12년; KO1_Education ∈ {4, 5}). 회귀모델에서 고학력군이 기준 범주였다. 단계별 순서형 점수화(저=1, 중=2, 고=3을 연속 추세 변수로 척도화)는 일차 위계 종점을 지원한다.")

add_heading("결과 변수 — 1년 급성악화", level=2)
add_para("일차 결과 변수는 베이스라인 12개월 이내 발생한 중등도 이상 급성악화로, KOCOSS 전향적 추적 기록(FU_Y1_M_exacerbation_YN)에서 확인하였다. 중등도 악화는 전신 스테로이드 및/또는 항생제 치료를, 중증 악화는 응급실 방문 또는 입원을 요한 경우로 정의하였다. Y1–Y3 확장 풀링 결과는 베이스라인 13~36개월 악화도 포함한다.")

add_heading("매개변수", level=2)
add_para("SGRQ 증상 영역을 사전에 일차 매개변수로 지정하였다. Pack-years, SGRQ 활동, SGRQ 영향 영역은 사전 비교 매개변수였다. 모든 매개변수는 베이스라인에서 측정되었다.")

add_heading("공변수", level=2)
add_para("모델은 연령(연속), 성별(남/여), 체질량지수(연속), BD 후 FEV1 % 예측치(연속), PRISm 상태(이항), 전년도 악화 이력(이항; KO1_P_EX_YN)을 보정하였다. 흡입 코르티코스테로이드 사용은 적응증에 의한 교란을 피하기 위해 공변수에서 제외하였다 [11].")

add_heading("통계 분석", level=2)
add_para("교육수준의 1년 악화 결과에 대한 총효과는 병원 군집화·교환가능 작업 상관구조의 GEE 로지스틱 회귀로 추정하였다 [23]. 세 개의 중첩 모델을 적합하였다: Model 1(비보정), Model 2(인구학적: 연령·성별·BMI), Model 3(완전 보정 — FEV1 % 예측·PRISm·전년도 악화 추가 포함). 일차 위계 종점은 단계별 순서형 추세(연속 교육 단계 점수화)였으며, 저-vs-고 이항 비교는 이차 종점이었다. 중-vs-고 비교는 기술적으로 보고하였다. Job=NaN은 결과·노출에 차등 결측이 없음을 근거(S8 참조)로 결측 무작위(MAR) 가정 하 일차 분석에 유지하였다. SGRQ 증상에서 MCAR 가정은 기각되었다(p=0.019).")
add_para("인과 매개분석은 VanderWeele 2014 4-way decomposition [20]을 따랐으며, Valeri & VanderWeele 2013 estimator [21]를 부트스트랩 2,000회로 적용하여 자연직접효과(NDE)·자연간접효과(NIE)·매개비율(PM)을 산출하였다. 순차적 무시가능성을 가정하였으며 §한계점에서 논의한다.")
add_para("민감도 분석(S1–S8, S6 제외)은 프로토콜에서 사전 지정되었다. 각 분석의 NIE OR, 95% CI, PM(추정 가능 시)을 표 4에 보고한다. 사전 지정되었던 S6(흡연자 한정 stratum)은 stratum 크기 부족(n=192)으로 인한 총효과 부트스트랩 분포 퇴화로 사후 제거하였다. 이 제외는 ICMJE 사전 지정-but-추정불가 분석 보고 지침에 따라 투명하게 보고한다.")
add_para("정량적 편향 평가. 두 가지 편향 정량을 보고하였다. 첫째, 간접효과에 대한 E-값[VanderWeele 2017; Linden 2020]을 산출하였다: 점추정치 1.49, 95% CI 하한 1.21이며, 이는 NIE를 무효화하는 데 필요한 미측정 매개변수–결과 교란인자의 강도를 의미한다. 둘째, 중간 정도의 미측정 교란 시나리오(R=1.10, 근사 ρ=0.05) 하 편향보정 NIE를 매개적 E-값 framework[Smith & VanderWeele 2019]에 따라 보수적 하한으로 산출하였다(보충 표 S3). 보정된 CI가 영가설을 통과하는 임계값(R≈1.21)을 명시 보고한다.")
add_para("코드 및 데이터 가용성. 분석 코드와 비식별화된 분석 데이터셋은 [DOI placeholder, 게재 승인 시점에 발급]에 보관된다. 전체 코드 저장소는 교신저자 소속기관 저장소에서 이용 가능하다. 모든 분석은 Python 3.11(statsmodels, numpy)과 R 4.3(gee, mediation, CMAverse)에서 수행되었다.")

add_image(IMG["Figure 1"], "그림 1. 개념적 방향성 비순환 그래프(DAG). 교육(노출) → 후보 매개변수(SGRQ 증상[일차, ★], SGRQ 활동, SGRQ 영향, 또는 pack-years[비교]) → 1년 급성악화(결과). 베이스라인 공변수(연령·성별·FEV1 % 예측·BMI·PRISm 상태·전년도 악화 이력)는 노출-매개·매개-결과 연관 모두를 보정한다. 베이스라인 측정·매개변수 평가·결과 확인의 시간 순서는 보충 그림 S1 참조.", width=5.5)

doc.add_page_break()

# =================== RESULTS ===================
add_heading("결과 (Results)", level=1)

add_heading("코호트 흐름 및 베이스라인 특성", level=2)
add_para("KOCOSS 분석 frame의 2,932명 등록자 중 220명이 임금 동질성 기준으로 제외되어(군인 8, 주부 124, 무직 67, 그 외 연령·기타 적격 기준 미충족) 최종 코호트 n=2,712를 구성하였다(그림 2). 완전 사례 매개분석 세트는 n=850이었으며, Y1–Y3 확장 풀링 분석 세트는 n=1,506이었다. 저학력군일수록 연령이 높고, 남성 비율이 높았으며, FEV1 % 예측치가 낮고, 전년도 악화 유병률이 높고, SGRQ 증상 점수가 높았다. 흡연력(pack-years)도 단계적 패턴을 보였으나(저 42.4 vs 고 33.0, p<0.001) 공식 매개분석에서는 교육-악화 연관을 매개하지 않았다(표 3 참조)(표 1).")

add_image(IMG["Figure 2"], "그림 2. 코호트 흐름도. KOCOSS 등록(n=2,932) → 임금 이질 직업군 제외(군인·주부·무직) → 최종 코호트(n=2,712) → 일차 매개분석 세트(n=850); 확장 Y1–Y3 분석 세트(n=1,506).", width=5.5)

# Table 1
add_para("표 1. 교육수준별 베이스라인 특성, KOCOSS 최종 코호트 (n=2,712; 교육 결측 59명 본 표 제외, 분석 코호트는 유지).", italic=True, size=10)
t1_h = ["변수", "저 (≤9년) n=1,170", "중 (10–12년) n=1,013", "고 (>12년) n=470", "p-값"]
t1_r = [
    ["연령(년) — 평균(SD)", "70.5 (7.3)", "67.5 (8.6)", "65.3 (9.9)", "<0.001"],
    ["남성, n(%)", "1,048 (89.6)", "940 (92.8)", "451 (96.0)", "<0.001"],
    ["BMI(kg/m²) — 평균(SD)", "23.4 (3.4)", "23.3 (3.3)", "23.3 (3.4)", "0.96"],
    ["FEV1 % 예측치 — 평균(SD)", "60.1 (18.8)", "62.1 (19.1)", "65.4 (18.5)", "<0.001"],
    ["PRISm, n(%)", "41 (3.5)", "41 (4.0)", "29 (6.2)", "0.049"],
    ["현재 흡연자, n(%)", "297 (25.4)", "302 (29.8)", "119 (25.4)", "0.043"],
    ["흡연력 (pack-years) — 평균(SD)", "42.4 (25.9)", "38.8 (24.2)", "33.0 (19.8)", "<0.001"],
    ["전년도 악화, n(%)", "199 (17.0)", "124 (12.2)", "52 (11.1)", "0.001"],
    ["SGRQ 증상 — 평균(SD)", "39.6 (21.9)", "34.8 (21.6)", "30.9 (19.9)", "<0.001"],
    ["SGRQ 활동 — 평균(SD)", "36.5 (26.4)", "32.0 (25.6)", "25.0 (24.3)", "<0.001"],
    ["SGRQ 영향 — 평균(SD)", "19.1 (21.0)", "15.4 (19.4)", "12.2 (15.8)", "<0.001"],
    ["Charlson 지수 — 중앙값(IQR)", "0.0 (0.0–1.0)", "0.0 (0.0–1.0)", "0.0 (0.0–1.0)", "0.077"],
]
add_table(t1_h, t1_r)
add_para("검정: 연속변수 = ANOVA; 범주변수 = chi-square. SGRQ 총점은 세 하위 영역 개별 보고로 대체.", italic=True, size=9)

doc.add_paragraph()

add_heading("교육수준의 1년 급성악화 총효과", level=2)
add_para("위계적 일차 종점에 따라, 단계별 순서형 추세 계수는 Model 3에서 통계적으로 유의하였다(단계별 OR 1.391, 95% CI 1.05~1.84, p-trend=0.020; Model 2 p-trend=0.014). 이차 이항 비교 — 저-vs-고 — 는 OR 1.71 (95% CI 0.77~3.82, p=0.191; 표 2)이었다. Y1–Y3 확장 풀링 결과 분석에서 저-vs-고 비교는 더 강하고 통계적으로 유의하였다(OR 1.81, 95% CI 1.09~3.00, p=0.023; 단계별 p-trend=0.008). 이는 더 넓은 추적 기간에서 단계적 노출-반응이 검출됨을 시사한다.")

# Table 2
add_para("표 2. 교육수준의 1년 급성악화 위험에 대한 보정된 총효과, KOCOSS 최종 코호트.", italic=True, size=10)
t2_h = ["비교", "Model 1 (비보정) OR (95% CI)", "Model 2 (인구학적) OR (95% CI)", "Model 3 (완전 보정) OR (95% CI)"]
t2_r = [
    ["저 vs 고 (이차)", "1.55 (0.81–2.97)", "1.62 (0.83–3.16)", "1.71 (0.77–3.82), p=0.191"],
    ["중 vs 고 (기술적)", "1.21 (0.61–2.40)", "1.25 (0.62–2.50)", "1.30 (0.59–2.86)"],
    ["단계별 추세 (일차 ★)", "1.36 (1.06–1.74)", "1.42 (1.07–1.88), p=0.014", "1.391 (1.05–1.84), p-trend=0.020"],
    ["Y1–Y3 풀링 — 저 vs 고", "—", "—", "1.81 (1.09–3.00), p=0.023; p-trend=0.008"],
]
add_table(t2_h, t2_r)

doc.add_paragraph()

add_heading("단일 매개변수 인과 매개분석", level=2)
add_para("일차 매개분석 세트(n=850)에서 사전 지정된 4개 매개변수를 단일 매개변수 모델에서 검정하였다. SGRQ 증상 영역만 통계적으로 유의한 매개변수였다(NIE OR 1.122, 95% CI 1.03~1.27; PM 21.9%; 표 3). Pack-years와 SGRQ 활동은 유의한 매개효과를 보이지 않았다. SGRQ 영향은 SGRQ 증상보다 작은 크기의 경계적 NIE를 기여하였다(1.06, 95% CI 1.00~1.14). α₁ 계수(4개 후보 매개변수에 대한 노출-매개변수 회귀)는 모두 양수였으며(+7.897, +5.270, +4.560, +7.596), 완전 보정 모델에서 교육-증상 연관의 예상된 방향을 확인하였다.")

# Table 3
add_para("표 3. 저-vs-고 교육수준의 1년 악화 위험에 대한 단일 매개변수 인과 매개분석 (n=850).", italic=True, size=10)
t3_h = ["매개변수", "α₁ (노출 → 매개)", "NIE OR (95% CI)", "NDE OR (95% CI)", "PM (%)", "유의 매개?"]
t3_r = [
    ["SGRQ 증상 (일차 ★)", "+7.897", "1.122 (1.03–1.27)", "1.61 (0.74–3.55)", "21.9", "예"],
    ["SGRQ 활동", "+5.270", "1.04 (0.97–1.13)", "1.66 (0.75–3.69)", "—", "아니오"],
    ["SGRQ 영향", "+4.560", "1.06 (1.00–1.14)", "1.62 (0.73–3.59)", "11.2", "경계"],
    ["Pack-years", "+7.596", "1.00 (0.98–1.04)", "1.71 (0.77–3.82)", "—", "아니오"],
]
add_table(t3_h, t3_r)

doc.add_paragraph()

add_heading("민감도 분석", level=2)
add_para("SGRQ 증상을 통한 일차 NIE는 보수적 편향보정 시나리오 R=1.10 하에서도 유의하게 유지되었다(보정 NIE OR 1.113, 95% CI 1.02~1.26; 표 4 'Bias-adjusted (R=1.10)' 행; 보충 표 S3). 보정된 CI 하한이 1을 통과하는 임계값은 R≈1.21이었다. 간접효과에 대한 E-값은 1.49(CI 하한 1.21)였다. 진행성 증상 질환에서의 역인과 편향에 가장 덜 취약한 GOLD 1–2 경증-중등증 하위 코호트(n=605)에서 방향이 보존되었다(NIE 1.197, 95% CI 1.06~1.39; PM 43.1%; 보충 표 S4). Y1–Y3 확장 풀링 분석은 더 강한 총효과를 보였다(OR 1.81, 95% CI 1.09~3.00). 직업 결측 93명을 제외한(S8) 분석은 OR 1.77 (0.78~4.02), 단계별 추세 1.39 (1.04~1.85, p=0.026)으로 MAR 하 일차 분석과 일치하였다. 매개분석 세트의 결측 역확률 가중치(S7, n=850)는 NIE OR 1.149 (1.05~1.29)로 일차보다 약간 강하였다. Charlson 추가 보정(S2, n=724)은 경계적 NIE 1.100 (1.00~1.24)을 산출하였다.")

# Table 4
add_para("표 4. 저-vs-고 교육수준의 1년 악화에 대한 SGRQ 증상-매개효과의 민감도 분석.", italic=True, size=10)
t4_h = ["분석", "n", "NIE OR (95% CI)", "TE / 저-고 OR (95% CI)", "PM (%)"]
t4_r = [
    ["일차 (Primary)", "850 (med set)", "1.122 (1.03–1.27) ★", "1.71 (0.77–3.82)", "21.9"],
    ["편향보정 (R=1.10)", "850", "1.113 (1.02–1.26) ★", "—", "—"],
    ["S1 FEV1 비보정", "1,415", "—", "1.82 (0.86–3.85)", "—"],
    ["S5 COPD 한정", "1,306", "—", "1.60 (0.78–3.27)", "—"],
    ["S2 + Charlson", "724", "1.100 (1.00–1.24) (경계)", "1.668 (0.94–3.60)", "18.6"],
    ["S7 결측 IPW", "850", "1.149 (1.05–1.29) ★", "1.878 (1.09–4.14)", "22.0"],
    ["Y1–Y3 풀링", "1,506", "—", "1.81 (1.09–3.00), p-trend=0.008", "—"],
    ["S8 Job=NaN 제외", "1,328", "—", "1.77 (0.78–4.02); 단계별 OR 1.39 (1.04–1.85), p=0.026", "—"],
    ["GOLD 1–2 하위 코호트 (보충 S4)", "605", "1.197 (1.06–1.39) ★", "1.531 (0.755–3.105)", "43.1"],
]
add_table(t4_h, t4_r)
add_para("★ = 95% CI가 1을 포함하지 않음.", italic=True, size=9)

add_image(IMG["Figure 3"], "그림 3. SGRQ 증상 영역을 통한 교육의 자연간접효과(NIE) — 일차·편향보정·7개 민감도 분석에 대한 forest plot (n=2,712 코호트). 가로 막대는 95% 신뢰구간, 수직 점선은 영가설(OR=1)을 표시한다. 일차 추정치는 ★로 표시하였다. S6(현재 흡연자 한정 stratum)은 stratum 크기 부족(n=192)으로 인한 총효과 부트스트랩 분포 퇴화로 사후 제거되었다 (방법 §민감도 참조).", width=6.0)

doc.add_page_break()

# =================== DISCUSSION ===================
add_heading("고찰 (Discussion)", level=1)

add_heading("주요 결과", level=2)
add_para("이 한국 다기관 전향적 COPD 또는 PRISm 코호트에서, 낮은 교육수준은 1년 급성악화 위험의 단계적·통계적으로 유의한 단계별 증가와 연관되었다(단계별 OR 1.391, p-trend=0.020). 효과는 Y1–Y3 확장 결과 기간에서 강화되었다(OR 1.81, p-trend=0.008). 공식 인과 매개분석은 SGRQ 증상 영역을 일차 매개 경로로 식별하였다: NIE OR 1.122 (95% CI 1.03~1.27), 총효과의 약 1/5을 설명한다. 간접효과는 보수적 교란보정 하에서도 유의하게 유지되었으며(R=1.10; 보정 NIE 1.113, 95% CI 1.02~1.26), 진행성 질환의 역인과 편향에 가장 덜 취약한 GOLD 1–2 경증-중등증 하위 코호트에서 방향이 보존되었다(NIE 1.197, 95% CI 1.06~1.39). Pack-years와 SGRQ 활동은 유의한 매개효과를 보이지 않았으며, SGRQ 영향은 더 작은 경계적 간접효과를 기여하였다. 이 패턴은 보편적 의료보장 체계 내에서 노출 강도나 활동 제한이 아닌 증상 경험이 교육과 악화 위험을 잇는 작동 경로임을 시사한다.")

add_heading("선행 문헌과의 비교 및 방향 삼각검증", level=2)
add_para("저학력 COPD 인구가 더 높은 악화 위험을 가진다는 본 연구의 방향적 발견은 덴마크 Glostrup [7], 캐나다 ICES [6, 8], 영국 CPRD 코호트 기반 관찰과 일관된다. 새로운 기여는 환자 인지 증상 부담을 통한 공식 매개 분해이다. 비중첩 미측정 교란 프로파일을 가진 출판된 코호트 — UK Biobank [30]와 ECLIPSE 코호트, 여기서 pack-years·동반질환·폐기능 보정이 다른 측정 도구를 사용 — 전반에 걸친 SES → COPD 결과 연관의 방향적 일관성은 단일 코호트 민감도 분석이 확립할 수 있는 수준을 넘어 인과 해석을 강화하는 외부 삼각검증을 제공한다.")

add_heading("가능한 기전", level=2)
add_para("SGRQ 증상 매개 경로의 기저에 세 가지 상호 배타적이지 않은 기전이 작동할 수 있다. 첫째, 낮은 교육수준은 건강 문해력 감소와 연관되며 [27, 28], 이는 증상 인지·도움요청을 지연시킬 수 있다. 둘째, 저학력 인구는 동반질환·약물순응도 변동성·진료 경로 탐색 어려움이 주도하는 차등 증상 증폭을 경험할 수 있으며, SGRQ 증상 영역은 이를 활동 제한이나 삶의 영향 영역보다 더 민감하게 포착한다 [15-18]. 셋째, 증상 보고 임계값 자체가 교육 단계별로 다르게 보정되어, 저학력 응답자가 더 낮은 임상 중증도에서 증상을 경험·보고하고 더 이른 악화 사건으로 전이될 수 있다.")

add_heading("강점", level=2)
add_para("주요 강점은 사전 임금 동질성 제외를 적용한 대규모 다기관 전향적 코호트(n=2,712, 44개 사이트), 노출·매개변수·공변수·1년 악화의 표준화된 확인, 부트스트랩 2,000회를 적용한 공식 4-way 인과 매개 framework이다. 견고성 전략은 편향보정 하한, E-값, 역인과를 다루는 GOLD 1–2 하위 코호트 민감도, 직업 결측에 대한 MAR 분석, 확장 Y1–Y3 풀링 결과 기간을 포함한다. S7 역확률 가중 분석(n=850)은 NIE OR 1.149 (1.05~1.29)를 산출하여 일차 추정치(1.122)보다 약간 강했다 — 간접효과가 완전 사례 선택의 산물이 아니라는 신뢰를 강화한다. S2 Charlson 추가 보정은 경계적 NIE 1.100 (1.00~1.24)을 산출하였으며, CI 하한 1.00은 잔여 동반질환 교란이 SGRQ 매개효과를 부분적으로 설명할 수 있음을 시사하지만 방향은 보존되었다.")

add_heading("한계점", level=2)
add_para("몇 가지 한계점을 고려해야 한다. 첫째, E-값 1.49는 간접효과가 중간 정도의 미측정 교란에 민감함을 나타낸다(R≈1.21 무효화 임계). KOCOSS의 가능한 미측정 요인에는 직업적 분진 이력, 거주지 PM2.5 노출, 가족 COPD 중증도가 포함되며 어느 것이든 이 임계에 근접할 수 있어, 편향보정 NIE(R=1.10)를 보수적 하한으로 보고한다. 둘째, 교육과 직업은 COPD 진단 후 베이스라인에서 확인되었으므로, GOLD 1–2 하위 코호트 견고성에도 불구하고 증상성 질환으로부터의 역인과를 완전히 배제할 수는 없다. 셋째, 직업이 93명에서 결측이었고(Job=NaN) MCAR 가정은 SGRQ 증상에서 기각되었지만(p=0.019), 결과·노출은 결측 관련 차등을 보이지 않았으며 S8 민감도 분석은 일차와 일치하였다. 분석은 따라서 MAR 가정 하에서 해석한다. 넷째, 인과 매개 framework은 경험적으로 검증할 수 없는 순차적 무시가능성 가정에 의존하며, 본 결과는 인과의 결정적 증명이 아닌 다수 가능 해석 중 하나로 제시된다. 다섯째, SGRQ 증상 영역과 본 연구의 악화 결과는 중첩되는 증상 구성개념을 가지지만, 만성 증상 부담과 치료 강화가 필요한 급성 사건의 시간적·질적·운영적 구분을 고려할 때 완전한 구성개념 등가성은 가능성이 낮다.")

add_heading("정책 및 인구 보건 함의", level=2)
add_para("보편적 의료보장 체계 내에서, 교육수준에 따른 잔여 1년 악화 gradient — 그 중 약 1/5이 환자 인지 증상 부담을 통해 흐르는 — 은 접근성 확대를 주된 lever에서 멀어지게 하며, 교육 자체를 증상 경험을 통해 작용하는 수정 가능한 상위 결정요인으로 식별한다. 교육 수준별 공중보건·진료체계 개입이 실행 가능 상위 경로로 따라온다: 저학력 COPD 계층에 만성질환 관리 프로그램 및 폐재활 우선 배정, 이 계층 외래 진료에 SGRQ 또는 COPD Assessment Test 같은 환자보고 증상 도구 통합, 저학력 인구의 증상 인지·도움요청 임계를 낮추는 증상 모니터링 진료 경로 개발. 그러한 차등 배정의 하류 실현가능성은 인구집단 수준 배치의 자원 강도를 낮추는 원격재활·디지털 자가관리 프로그램의 새로운 증거 [Cox 2022; Gloeckl 2025; Bourne 2017]에 의해 강화되었다. 본 연구에서 pack-years 매개효과의 부재는 기존 금연 정책을 대체해야 함을 의미하지 않는다. 교육 수준별 증상 모니터링 진료가 그것들을 대체하기보다는 함께 추가되어야 한다. 본 함의는 가설 생성적이며 구현 연구에서 확인이 필요하다. 그럼에도 불구하고 저학력 COPD 인구가 어떻게 증상을 경험하고 보고하는지 — 단지 의료에 도달하는지 여부가 아니라 — 를 다루는 것이 보편적 의료보장 체계에서 더 높은 산출의 인구 lever일 수 있다.")

add_heading("결론", level=2)
add_para("한국 다기관 전향적 코호트에서, 1년 COPD 또는 PRISm 악화 위험의 단계적 교육수준 gradient를 식별하였으며, 그 중 약 1/5이 SGRQ 증상 영역으로 측정된 환자 인지 증상 부담을 통해 매개되었다. 간접효과는 보수적 미측정 교란보정 하에서 유의하게 유지되었으며 경증-중등증 하위 코호트에서 방향이 보존되었다. 본 결과는 교육수준을 증상 경험을 통해 작용하는 수정 가능한 상위 결정요인으로 식별하며, 보편적 의료보장 COPD 체계 내 상위 정책 lever로서 교육 수준별 증상 모니터링 진료 경로 및 우선 자원 배정을 지지한다.")

doc.add_page_break()

# =================== SUPPLEMENTARY ===================
add_heading("보충 자료 (Supplementary material)", level=1)
add_para("보충 그림 S1 — KOCOSS의 베이스라인·매개변수·결과 측정 시간 순서(아래). 보충 그림 S2 — 추적 기간 window sweep(Y1~Y5)에 따른 총효과(TE) 및 자연간접효과(NIE) 시간 robustness. 보충 표 S3 — 정량적 편향 평가(E-값 sweep, ρ ∈ {0.05, 0.10, 0.15, 0.20, 0.25}); 보충 표 S4 — GOLD 1–2 하위 코호트 민감도(n=605); 보충 표 S8 — Job-NaN MAR 검정. 이 표들은 별도 파일(S3_ignorability_sensitivity, S4_gold12_subcohort, S8_mcar_test)로 JSON 및 Markdown 형식으로 제공된다.")
add_image(IMG["Supp S1"], "보충 그림 S1. KOCOSS 코호트에서 베이스라인 측정(T0), 매개변수 평가, 결과 확인(일차 T+12개월; 확장 Y1–Y3 풀링 분석 T+12~T+36개월)의 시간 순서. 교육과 직업은 베이스라인에서 확인되었다. 진행성 증상 질환으로부터의 역인과는 완전히 배제할 수 없으나 GOLD 1–2 하위 코호트 민감도 분석(표 4 / 보충 표 S4)에서 다루어진다.", width=6.0)
add_image(IMG["Supp S2"], "보충 그림 S2. 교육수준(저 vs 고)의 5개 누적 추적 window(Y1 일차~Y1–Y5)에 걸친 총효과(TE) 및 SGRQ 증상 통한 자연간접효과(NIE) follow-up window sweep. 좌측 패널: TE 오즈비(로그 척도); 우측 패널: NIE 오즈비(선형 척도). 95% 신뢰구간 포함. NIE 점추정치는 전 window에 걸쳐 안정(범위 1.056–1.122)하며, TE는 Y1–Y3에서 통계적 유의성 도달(OR 1.81, 95% CI 1.09–3.00). 출처: followup_sweep.json.", width=6.0)

doc.add_page_break()

# =================== DECLARATIONS ===================
add_heading("선언 (Declarations)", level=1)

decls_kr = [
    ("연구비 지원 (Funding).", "본 연구는 대한민국 질병관리청 국립보건연구원 연구사업의 지원을 받아 수행되었다(과제번호 2016ER670100, 2016ER670101, 2016ER670102, 2018ER670100, 2018ER670101, 2018ER670102, 2021ER120500, 2021ER120501, 2021ER120502, 2024ER120100, 2024ER120101, 2024ER120102, 2023NI00702)."),
    ("이해상충 (Competing interests).", "저자들은 이해상충이 없음을 선언한다."),
    ("연구윤리 승인 및 동의 (Ethics approval).", "한국 COPD 코호트 연구(KOCOSS)에 참여한 모든 병원은 해당 임상시험심사위원회의 승인을 받았으며, 건국대학교병원(IRB No. KHH1010338)을 포함한다. 모든 참가자는 서면 동의서를 제공하였다."),
    ("출판 동의 (Consent for publication).", "해당 없음."),
    ("데이터 가용성 (Data availability).", "분석 코드와 비식별화된 분석 데이터셋은 동료심사 전 인용 가능한 DOI와 함께 Zenodo에 보관될 예정이다(DOI placeholder — 제출 시점 삽입; 방법 §통계 분석 참조). 원시 KOCOSS 데이터는 KOCOSS 데이터 공유 정책 및 한국 관련 규정의 적용을 받으며, 요청은 교신저자가 검토한다."),
    ("감사의 글 (Acknowledgements).", "한국 COPD 코호트 연구(KOCOSS)의 참가자와 연구진에게 감사드린다."),
    ("저자 기여 (Author contributions).", "Eunjin Kwon: 개념화, 방법론, 데이터 큐레이션, 형식적 분석, 시각화, 검증, 원고 초안 작성, 검토 및 편집. Won Seo Yoon: 방법론, 데이터 큐레이션, 검증, 검토 및 편집. Ji-Yong Moon: 조사, 검토 및 편집. Gi Ho Lee: 데이터 큐레이션, 검토 및 편집. Yong-Il Hwang: 조사, 데이터 큐레이션, 검토 및 편집. Kwang-Ha Yoo: 조사, 감독, 검토 및 편집. Young-Youl Kim: 개념화, 감독, 연구비 확보, 프로젝트 관리, 검토 및 편집. Youlim Kim: 개념화, 감독, 프로젝트 관리, 검토 및 편집."),
]
for label, text in decls_kr:
    add_para_bold_label(label, text)

note = doc.add_paragraph()
r = note.add_run("[제출 전 확인 필요] 연구비/연구윤리/이해상충은 자매 논문(Kwon E et al. 2026 Respir Res) 기준. 저자 기여는 자매 논문의 표준 CRediT 패턴을 준용. 제1저자가 본 분석의 실제 기여도에 따라 확정·수정 필요.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x7a, 0x7a, 0x7a); set_kr_font(r)

doc.add_page_break()

# =================== REFERENCES ===================
add_heading("References", level=1)
add_para("(참고문헌은 영문 원문을 그대로 유지함. ERJ 제출 표준)", italic=True, size=10)
doc.add_paragraph()

refs = [
    "MacLeod M, Papi A, Contoli M, et al. Chronic obstructive pulmonary disease exacerbation fundamentals: diagnosis, treatment, prevention and disease impact. Respirology 2021;26:532–551. doi:10.1111/resp.14041",
    "Global Initiative for Chronic Obstructive Lung Disease. Global Strategy for Prevention, Diagnosis and Management of COPD: 2025 Report. [verify — goldcopd.org]",
    "GBD 2019 Chronic Respiratory Diseases Collaborators. Prevalence and attributable health burden of chronic respiratory diseases, 1990–2017. Lancet Respir Med 2020;8:585–596. doi:10.1016/S2213-2600(20)30105-3",
    "Yang IA, Jenkins CR, Salvi SS. Chronic obstructive pulmonary disease in never-smokers: risk factors, pathogenesis, and implications. Lancet Respir Med 2022;10:497–511. doi:10.1016/S2213-2600(21)00506-3",
    "Stolz D, Mkorombindo T, Schumann DM, et al. Towards the elimination of chronic obstructive pulmonary disease: a Lancet Commission. Lancet 2022;400:921–972. doi:10.1016/S0140-6736(22)01273-9",
    "Gershon AS, Dolmage TE, Stephenson A, Jackson B. COPD and socioeconomic status: a systematic review. COPD 2012;9:216–226. doi:10.3109/15412555.2011.648030 [verify]",
    "Prescott E, Lange P, Vestbo J. Socioeconomic status, lung function and admission to hospital for COPD. Eur Respir J 1999;13:1109–1114. [verify]",
    "Marmot M. Social determinants of health inequalities. Lancet 2005;365:1099–1104. doi:10.1016/S0140-6736(05)71146-6",
    "Galobardes B, Shaw M, Lawlor DA, Lynch JW, Davey Smith G. Indicators of socioeconomic position (part 1). J Epidemiol Community Health 2006;60:7–12. doi:10.1136/jech.2004.023531",
    "Wan ES, Fortis S, Regan EA, et al. Epidemiology, genetics, and subtyping of preserved ratio impaired spirometry (PRISm) in COPDGene. Respir Res 2014;15:89. [verify]",
    "Bhatt SP, Abadi E, Anzueto A, et al. A multidimensional diagnostic approach for COPD. JAMA 2025;333:2164–2175. doi:10.1001/jama.2025.7358",
    "Hurst JR, Vestbo J, Anzueto A, et al. Susceptibility to exacerbation in COPD. N Engl J Med 2010;363:1128–1138. doi:10.1056/NEJMoa0909883",
    "Celli BR, Fabbri LM, Aaron SD, et al. An updated definition and severity classification of COPD exacerbations: the Rome Proposal. Am J Respir Crit Care Med 2021;204:1251–1258. doi:10.1164/rccm.202108-1819PP",
    "Divo M, Cote C, de Torres JP, et al. Comorbidities and risk of mortality in COPD. Am J Respir Crit Care Med 2012;186:155–161. doi:10.1164/rccm.201201-0034OC",
    "Jones PW, Quirk FH, Baveystock CM, Littlejohns P. A self-complete measure of health status for chronic airflow limitation: the SGRQ. Am Rev Respir Dis 1992;145:1321–1327.",
    "Jones PW. SGRQ: MCID. COPD 2005;2:75–79.",
    "Kessler R, Partridge MR, Miravitlles M, et al. Symptom variability in patients with severe COPD. Eur Respir J 2011;37:264–272.",
    "Choi JY, Yoon HK, Shin KC, et al. CAT score and SGRQ definitions of chronic bronchitis. Int J Chron Obstruct Pulmon Dis 2019;14:3043–3052.",
    "Jang JG, Kim Y, Lee JK, et al. Clinical characteristics of individuals with COPD, pre-COPD and smokers with normal lung function in Korea. Tuberc Respir Dis (Seoul) 2025;89:75–85.",
    "VanderWeele TJ. A unification of mediation and interaction: a 4-way decomposition. Epidemiology 2014;25:749–761.",
    "Valeri L, VanderWeele TJ. Mediation analysis allowing for exposure–mediator interactions and causal interpretation. Psychol Methods 2013;18:137–150.",
    "Imai K, Keele L, Tingley D. A general approach to causal mediation analysis. Psychol Methods 2010;15:309–334.",
    "Zeger SL, Liang KY. Longitudinal data analysis for discrete and continuous outcomes. Biometrics 1986;42:121–130.",
    "Seaman SR, White IR. Review of inverse probability weighting for dealing with missing data. Stat Methods Med Res 2013;22:278–295.",
    "VanderWeele TJ, Ding P. Sensitivity analysis in observational research: introducing the E-value. Ann Intern Med 2017;167:268–274.",
    "Linden A, Mathur MB, VanderWeele TJ. Conducting sensitivity analysis for unmeasured confounding using E-values. Stata J 2020;20:162–175.",
    "Nutbeam D. The evolving concept of health literacy. Soc Sci Med 2008;67:2072–2078.",
    "Berkman ND, Sheridan SL, Donahue KE, Halpern DJ, Crotty K. Low health literacy and health outcomes. Ann Intern Med 2011;155:97–107.",
    "Effing TW, Vercoulen JH, Bourbeau J, et al. Definition of a COPD self-management intervention. Eur Respir J 2016;48:46–54.",
    "Han YY, Yan Q, Chen W, Celedón JC. Child maltreatment, anxiety and depression, and asthma among British adults in the UK Biobank. Eur Respir J 2022;60(4):2103160.",
    "Jones PW, Rutten-van Mölken MPMH, Agusti A, et al. Reporting patient-reported outcomes in COPD. Eur Respir J 2019;54:1900168.",
    "Lamberton CE, Mosher CL. Review of the evidence for pulmonary rehabilitation in COPD. Respir Care 2024;69:686–696.",
    "Spruit MA, Singh SJ, Garvey C, et al. ATS/ERS statement: key concepts and advances in pulmonary rehabilitation. Am J Respir Crit Care Med 2013;188:e13–e64.",
    "Cox NS, McDonald CF, Mahal A, et al. Telerehabilitation for chronic respiratory disease: a randomised controlled equivalence trial. Thorax 2022;77:643–651.",
    "Gloeckl R, Spielmanns M, Stankeviciene A, et al. Smartphone application-based pulmonary rehabilitation in COPD. Thorax 2025;80:209–217.",
    "Bourne S, DeVos R, North M, et al. Online versus face-to-face pulmonary rehabilitation for patients with COPD. BMJ Open 2017;7:e014580.",
    "Kwon E, Yoon WS, Moon JY, Lee GH, Hwang YI, Yoo KH, Kim YY, Kim Y. Cardiovascular components of the COTE index predict acute exacerbations and healthcare costs in patients with chronic obstructive pulmonary disease: a nationwide linked cohort study. Respir Res 2026;27. doi:10.1186/s12931-026-03677-4",
    "Smith LH, VanderWeele TJ. Mediational E-values: approximate sensitivity analysis for unmeasured mediator–outcome confounding. Epidemiology 2019;30:835–837. doi:10.1097/EDE.0000000000001064",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. "); r.bold = True; r.font.size = Pt(10); set_kr_font(r)
    r2 = p.add_run(ref); r2.font.size = Pt(10); set_kr_font(r2)

doc.save(OUT)
print(f"DONE: {OUT}")
print(f"Size: {OUT.stat().st_size} bytes")
print(f"Refs: {len(refs)}")
